#!/usr/bin/env python3
"""Convert Markdown into Word, Excel, or PDF files.

This helper intentionally keeps parsing conservative. It handles common AI
Markdown well: headings, paragraphs, lists, blockquotes, fenced code, and pipe
tables. For higher-fidelity PDF/Word output, the surrounding skill should prefer
Pandoc when it is available.
"""

from __future__ import annotations

import argparse
import os
import re
import shutil
import subprocess
import sys
import tempfile
from dataclasses import dataclass
from html import escape
from pathlib import Path
from typing import Iterable, List, Sequence


@dataclass
class Block:
    kind: str
    text: str = ""
    level: int = 0
    rows: List[List[str]] | None = None


@dataclass
class PdfFontChoice:
    regular: str
    bold: str
    source: str


@dataclass
class MarkdownFeature:
    key: str
    label: str
    native_behavior: str


TABLE_SEPARATOR_RE = re.compile(r"^\s*\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?\s*$")

PDF_FONT_CANDIDATES = [
    ("Microsoft YaHei", r"C:\Windows\Fonts\msyh.ttc", r"C:\Windows\Fonts\msyhbd.ttc"),
    ("Noto Sans SC", r"C:\Windows\Fonts\NotoSansSC-VF.ttf", r"C:\Windows\Fonts\NotoSansSC-VF.ttf"),
    ("SimHei", r"C:\Windows\Fonts\simhei.ttf", r"C:\Windows\Fonts\simhei.ttf"),
    ("SimSun", r"C:\Windows\Fonts\simsun.ttc", r"C:\Windows\Fonts\simsunb.ttf"),
    ("DengXian", r"C:\Windows\Fonts\Deng.ttf", r"C:\Windows\Fonts\Dengb.ttf"),
    ("PingFang SC", "/System/Library/Fonts/PingFang.ttc", "/System/Library/Fonts/PingFang.ttc"),
    ("Hiragino Sans GB", "/System/Library/Fonts/Hiragino Sans GB.ttc", "/System/Library/Fonts/Hiragino Sans GB.ttc"),
    ("Noto Sans CJK", "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc", "/usr/share/fonts/opentype/noto/NotoSansCJK-Bold.ttc"),
    ("Noto Sans CJK", "/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc", "/usr/share/fonts/truetype/noto/NotoSansCJK-Bold.ttc"),
    ("Source Han Sans", "/usr/share/fonts/opentype/source-han-sans/SourceHanSansCN-Regular.otf", "/usr/share/fonts/opentype/source-han-sans/SourceHanSansCN-Bold.otf"),
]

PANDOC_MARKDOWN_FORMAT = (
    "markdown"
    "+tex_math_dollars"
    "+tex_math_single_backslash"
    "+pipe_tables"
    "+grid_tables"
    "+footnotes"
    "+task_lists"
    "+strikeout"
    "+definition_lists"
    "+yaml_metadata_block"
    "+raw_html"
)

MARKDOWN_READER_CSS = """
:root {
  color-scheme: light;
  --fg: #24292f;
  --muted: #57606a;
  --border: #d0d7de;
  --subtle: #f6f8fa;
  --accent: #0969da;
}
@page {
  margin: 18mm 16mm;
}
html {
  background: #fff;
}
body {
  box-sizing: border-box;
  max-width: 920px;
  margin: 0 auto;
  padding: 32px 28px 48px;
  color: var(--fg);
  font-family: "Microsoft YaHei", "Noto Sans CJK SC", "PingFang SC", "Helvetica Neue", Arial, sans-serif;
  font-size: 16px;
  line-height: 1.65;
}
h1, h2, h3, h4, h5, h6 {
  margin: 24px 0 12px;
  font-weight: 700;
  line-height: 1.25;
}
h1 {
  margin-top: 0;
  padding-bottom: .3em;
  font-size: 2em;
  border-bottom: 1px solid var(--border);
}
h2 {
  padding-bottom: .25em;
  font-size: 1.5em;
  border-bottom: 1px solid var(--border);
}
h3 { font-size: 1.25em; }
p, blockquote, ul, ol, dl, table, pre, details {
  margin-top: 0;
  margin-bottom: 16px;
}
a { color: var(--accent); text-decoration: none; }
blockquote {
  padding: 0 1em;
  color: var(--muted);
  border-left: .25em solid var(--border);
}
ul, ol { padding-left: 2em; }
li + li { margin-top: .25em; }
input[type="checkbox"] {
  margin-right: .45em;
}
code, kbd, samp {
  font-family: ui-monospace, SFMono-Regular, Consolas, "Liberation Mono", Menlo, monospace;
  font-size: 85%;
}
code {
  padding: .2em .4em;
  background-color: rgba(175,184,193,.2);
  border-radius: 6px;
}
pre {
  padding: 16px;
  overflow: auto;
  background-color: var(--subtle);
  border-radius: 6px;
}
pre code {
  padding: 0;
  background: transparent;
  border-radius: 0;
  font-size: 100%;
}
table {
  display: table;
  width: 100%;
  border-spacing: 0;
  border-collapse: collapse;
}
td, th {
  padding: 6px 13px;
  border: 1px solid var(--border);
  vertical-align: top;
}
th {
  font-weight: 600;
  background: #f6f8fa;
}
tr:nth-child(2n) {
  background: #f6f8fa;
}
mark {
  padding: .1em .25em;
  background: #fff8c5;
  border-radius: 4px;
}
math[display="block"] {
  display: block;
  margin: 1em auto;
  text-align: center;
  font-size: 1.12em;
}
.footnotes {
  margin-top: 32px;
  padding-top: 16px;
  color: var(--muted);
  border-top: 1px solid var(--border);
  font-size: .9em;
}
.sourceCode,
div.sourceCode {
  background: transparent;
}
.mermaid-unrendered {
  position: relative;
  border: 1px dashed var(--border);
}
.mermaid-unrendered::before {
  content: "Mermaid source";
  display: block;
  margin-bottom: 8px;
  color: var(--muted);
  font-family: "Microsoft YaHei", Arial, sans-serif;
  font-size: 12px;
}
@media print {
  body {
    max-width: none;
    padding: 0;
    font-size: 14px;
  }
  h1, h2 {
    break-after: avoid;
  }
  pre, table, blockquote {
    break-inside: avoid;
  }
}
"""


def strip_inline_markup(text: str) -> str:
    text = re.sub(r"!\[([^\]]*)\]\([^)]+\)", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    text = re.sub(r"(`+)(.*?)\1", r"\2", text)
    text = re.sub(r"(\*\*|__)(.*?)\1", r"\2", text)
    text = re.sub(r"(\*|_)(.*?)\1", r"\2", text)
    return text.replace("\\|", "|").strip()


def detect_markdown_features(markdown: str) -> List[MarkdownFeature]:
    feature_specs = [
        (
            "math",
            "TeX/LaTeX math",
            "native preserves formula source as text; use Pandoc for rendered equations",
            [r"(?s)(?<!\\)\$\$.+?(?<!\\)\$\$", r"(?<!\\)\$(?!\s|\$).+?(?<!\s|\\)\$", r"\\\(.+?\\\)", r"(?s)\\\[.+?\\\]"],
        ),
        (
            "footnotes",
            "footnotes",
            "native keeps footnote markers as text; use Pandoc for real footnotes",
            [r"(?m)^\[\^[^\]]+\]:", r"\[\^[^\]]+\]"],
        ),
        (
            "raw-html",
            "raw HTML",
            "native may simplify or preserve visible text only; use Pandoc/HTML pipeline for layout",
            [r"(?m)^\s*</?[A-Za-z][\w:-]*(?:\s|>|/>)", r"<(table|div|span|sup|sub|kbd|mark|details|summary)\b"],
        ),
        (
            "diagram",
            "diagram fence",
            "native preserves diagram source as a code block; render Mermaid/PlantUML first for images",
            [r"(?m)^```(?:mermaid|plantuml|graphviz|dot|vega-lite|chart)\b"],
        ),
        (
            "task-list",
            "task lists",
            "native keeps checkbox markers as text; Pandoc can map them better",
            [r"(?m)^\s*[-*+]\s+\[[ xX]\]\s+"],
        ),
        (
            "definition-list",
            "definition lists",
            "native treats definitions as paragraphs; Pandoc preserves list structure",
            [r"(?m)^\s*:\s+\S"],
        ),
        (
            "grid-table",
            "grid tables",
            "native supports pipe tables only; Pandoc handles grid tables",
            [r"(?m)^\+[=-]+(?:\+[=-]+)+\+$"],
        ),
        (
            "citation",
            "citations",
            "native keeps citation keys as text; Pandoc can process citations with bibliography data",
            [r"\[@[A-Za-z0-9_:-]+"],
        ),
        (
            "admonition",
            "admonitions/directives",
            "native treats admonitions as text or code; use Pandoc filters or HTML pipeline for boxes",
            [r"(?m)^\s*(:::+|!!!\s+\w+)"],
        ),
        (
            "image",
            "images",
            "native keeps image alt text; use Pandoc/HTML pipeline for embedded images",
            [r"!\[[^\]]*\]\([^)]+\)"],
        ),
        (
            "extended-inline",
            "extended inline styles",
            "native simplifies formatting; Pandoc preserves more inline semantics",
            [r"~~[^~]+~~", r"==[^=]+=="],
        ),
    ]

    found: List[MarkdownFeature] = []
    seen = set()
    for key, label, native_behavior, patterns in feature_specs:
        for pattern in patterns:
            if re.search(pattern, markdown):
                if key not in seen:
                    found.append(MarkdownFeature(key, label, native_behavior))
                    seen.add(key)
                break
    return found


def warn_about_native_limitations(features: Sequence[MarkdownFeature], target: str) -> None:
    if not features:
        return
    feature_labels = ", ".join(feature.label for feature in features)
    print(
        f"Warning: detected complex Markdown for {target}: {feature_labels}. "
        "Using native fallback, so some features may be preserved as plain text.",
        file=sys.stderr,
        flush=True,
    )
    for feature in features:
        print(f"  - {feature.label}: {feature.native_behavior}", file=sys.stderr, flush=True)


def split_table_row(line: str) -> List[str]:
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    cells: List[str] = []
    current: List[str] = []
    escaped = False
    for char in line:
        if escaped:
            current.append(char)
            escaped = False
        elif char == "\\":
            escaped = True
        elif char == "|":
            cells.append(strip_inline_markup("".join(current)))
            current = []
        else:
            current.append(char)
    cells.append(strip_inline_markup("".join(current)))
    return cells


def is_table_start(lines: Sequence[str], index: int) -> bool:
    return (
        index + 1 < len(lines)
        and "|" in lines[index]
        and TABLE_SEPARATOR_RE.match(lines[index + 1]) is not None
    )


def parse_markdown(markdown: str) -> List[Block]:
    lines = markdown.splitlines()
    blocks: List[Block] = []
    paragraph: List[str] = []
    in_code = False
    code_lines: List[str] = []
    index = 0

    def flush_paragraph() -> None:
        nonlocal paragraph
        if paragraph:
            blocks.append(Block(kind="paragraph", text=strip_inline_markup(" ".join(paragraph))))
            paragraph = []

    while index < len(lines):
        line = lines[index]
        stripped = line.strip()

        if stripped.startswith("```") or stripped.startswith("~~~"):
            if in_code:
                blocks.append(Block(kind="code", text="\n".join(code_lines)))
                code_lines = []
                in_code = False
            else:
                flush_paragraph()
                in_code = True
            index += 1
            continue

        if in_code:
            code_lines.append(line)
            index += 1
            continue

        if not stripped:
            flush_paragraph()
            index += 1
            continue

        if is_table_start(lines, index):
            flush_paragraph()
            rows = [split_table_row(lines[index])]
            index += 2
            while index < len(lines) and "|" in lines[index] and lines[index].strip():
                rows.append(split_table_row(lines[index]))
                index += 1
            blocks.append(Block(kind="table", rows=rows))
            continue

        heading = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if heading:
            flush_paragraph()
            blocks.append(
                Block(
                    kind="heading",
                    level=len(heading.group(1)),
                    text=strip_inline_markup(heading.group(2)),
                )
            )
            index += 1
            continue

        quote = re.match(r"^>\s?(.*)$", stripped)
        if quote:
            flush_paragraph()
            blocks.append(Block(kind="quote", text=strip_inline_markup(quote.group(1))))
            index += 1
            continue

        unordered = re.match(r"^[-*+]\s+(.+)$", stripped)
        ordered = re.match(r"^\d+[.)]\s+(.+)$", stripped)
        if unordered or ordered:
            flush_paragraph()
            match = unordered or ordered
            blocks.append(
                Block(
                    kind="bullet" if unordered else "number",
                    text=strip_inline_markup(match.group(1)),
                )
            )
            index += 1
            continue

        paragraph.append(stripped)
        index += 1

    flush_paragraph()
    if code_lines:
        blocks.append(Block(kind="code", text="\n".join(code_lines)))
    return blocks


def ensure_parent(path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)


def paragraph_text(text: str) -> str:
    return escape(text).replace("\n", "<br/>")


def convert_docx(blocks: Sequence[Block], output: Path) -> None:
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError as exc:
        raise SystemExit("Missing dependency for Word output: install python-docx.") from exc

    document = Document()
    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)

    for block in blocks:
        if block.kind == "heading":
            document.add_heading(block.text, level=min(block.level, 4))
        elif block.kind == "bullet":
            document.add_paragraph(block.text, style="List Bullet")
        elif block.kind == "number":
            document.add_paragraph(block.text, style="List Number")
        elif block.kind == "quote":
            paragraph = document.add_paragraph()
            paragraph.style = "Intense Quote"
            paragraph.add_run(block.text)
        elif block.kind == "code":
            paragraph = document.add_paragraph()
            run = paragraph.add_run(block.text)
            run.font.name = "Consolas"
            run.font.size = Pt(9)
        elif block.kind == "table" and block.rows:
            width = max(len(row) for row in block.rows)
            table = document.add_table(rows=len(block.rows), cols=width)
            table.style = "Table Grid"
            for row_index, row in enumerate(block.rows):
                for col_index in range(width):
                    table.cell(row_index, col_index).text = row[col_index] if col_index < len(row) else ""
        elif block.text:
            document.add_paragraph(block.text)

    ensure_parent(output)
    document.save(output)


def iter_tables(blocks: Iterable[Block]) -> Iterable[List[List[str]]]:
    for block in blocks:
        if block.kind == "table" and block.rows:
            yield block.rows


def convert_xlsx(blocks: Sequence[Block], output: Path) -> int:
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill
        from openpyxl.utils import get_column_letter
    except ImportError as exc:
        raise SystemExit("Missing dependency for Excel output: install openpyxl.") from exc

    workbook = Workbook()
    default_sheet = workbook.active
    workbook.remove(default_sheet)

    table_count = 0
    for table_count, rows in enumerate(iter_tables(blocks), start=1):
        sheet = workbook.create_sheet(f"Table {table_count}")
        for row in rows:
            sheet.append(row)
        for cell in sheet[1]:
            cell.font = Font(bold=True)
            cell.fill = PatternFill("solid", fgColor="D9EAF7")
        for column_cells in sheet.columns:
            max_length = max(len(str(cell.value or "")) for cell in column_cells)
            sheet.column_dimensions[get_column_letter(column_cells[0].column)].width = min(max(max_length + 2, 10), 48)

    if table_count == 0:
        sheet = workbook.create_sheet("Markdown")
        sheet.append(["Type", "Level", "Content"])
        for block in blocks:
            if block.kind != "table":
                sheet.append([block.kind, block.level or "", block.text])
        for cell in sheet[1]:
            cell.font = Font(bold=True)
            cell.fill = PatternFill("solid", fgColor="D9EAF7")
        sheet.column_dimensions["A"].width = 16
        sheet.column_dimensions["B"].width = 10
        sheet.column_dimensions["C"].width = 80

    ensure_parent(output)
    workbook.save(output)
    return table_count


def register_pdf_fonts(font_path: str | None = None, bold_font_path: str | None = None) -> PdfFontChoice:
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont
    from reportlab.pdfbase.ttfonts import TTFont

    custom_regular = font_path or os.environ.get("MARK_TO_WIN_PDF_FONT")
    custom_bold = bold_font_path or os.environ.get("MARK_TO_WIN_PDF_BOLD_FONT")
    candidates = []
    if custom_regular:
        candidates.append(("Custom", custom_regular, custom_bold or custom_regular))
    candidates.extend(PDF_FONT_CANDIDATES)

    for label, regular_path, bold_path in candidates:
        regular_file = Path(regular_path).expanduser()
        if not regular_file.exists():
            continue
        try:
            pdfmetrics.registerFont(TTFont("MarkToWinCJK", str(regular_file)))
            selected_bold = "MarkToWinCJK"
            bold_source = str(regular_file)
            if bold_path:
                bold_file = Path(bold_path).expanduser()
                if bold_file.exists():
                    try:
                        pdfmetrics.registerFont(TTFont("MarkToWinCJKBold", str(bold_file)))
                        selected_bold = "MarkToWinCJKBold"
                        bold_source = str(bold_file)
                    except Exception:
                        selected_bold = "MarkToWinCJK"
            source = f"{label}: {regular_file}"
            if selected_bold != "MarkToWinCJK":
                source = f"{source}; bold: {bold_source}"
            return PdfFontChoice("MarkToWinCJK", selected_bold, source)
        except Exception:
            continue

    try:
        pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
        return PdfFontChoice("STSong-Light", "STSong-Light", "ReportLab STSong-Light CID fallback")
    except Exception:
        return PdfFontChoice("Helvetica", "Helvetica-Bold", "Helvetica fallback; CJK text may not render")


def table_column_widths(rows: Sequence[Sequence[str]], available_width: float) -> List[float]:
    width = max(len(row) for row in rows)
    weights: List[int] = []
    for col_index in range(width):
        max_len = max((len(row[col_index]) if col_index < len(row) else 0) for row in rows)
        weights.append(max(max_len, 6))
    total = sum(weights) or width
    min_width = min(72, available_width / max(width, 1))
    widths = [max(min_width, available_width * weight / total) for weight in weights]
    scale = available_width / sum(widths)
    return [item * scale for item in widths]


def normalize_table_rows(rows: Sequence[Sequence[str]]) -> List[List[str]]:
    width = max(len(row) for row in rows)
    return [[row[index] if index < len(row) else "" for index in range(width)] for row in rows]


def convert_pdf(blocks: Sequence[Block], output: Path, font_path: str | None = None, bold_font_path: str | None = None) -> PdfFontChoice:
    try:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.platypus import Paragraph, Preformatted, SimpleDocTemplate, Spacer, Table, TableStyle
    except ImportError as exc:
        raise SystemExit("Missing dependency for PDF output: install reportlab, or use pandoc.") from exc

    font_choice = register_pdf_fonts(font_path, bold_font_path)
    styles = getSampleStyleSheet()
    styles["Normal"].fontName = font_choice.regular
    styles["BodyText"].fontName = font_choice.regular
    styles["BodyText"].fontSize = 10.5
    styles["BodyText"].leading = 16
    styles["Heading1"].fontName = font_choice.bold
    styles["Heading1"].fontSize = 22
    styles["Heading1"].leading = 28
    styles["Heading1"].spaceAfter = 14
    styles["Heading2"].fontName = font_choice.bold
    styles["Heading2"].fontSize = 15
    styles["Heading2"].leading = 21
    styles["Heading2"].spaceBefore = 10
    styles["Heading2"].spaceAfter = 8
    quote_style = ParagraphStyle(
        "MarkToWinQuote",
        parent=styles["BodyText"],
        leftIndent=12,
        textColor=colors.HexColor("#334155"),
        borderColor=colors.HexColor("#CBD5E1"),
        borderWidth=0,
        borderPadding=0,
    )
    code_style = ParagraphStyle(
        "MarkToWinCode",
        parent=styles["Code"],
        fontName=font_choice.regular,
        fontSize=8.5,
        leading=12,
        leftIndent=10,
        rightIndent=10,
        backColor=colors.HexColor("#F6F8FA"),
        textColor=colors.HexColor("#24292F"),
    )
    cell_style = ParagraphStyle(
        "MarkToWinTableCell",
        parent=styles["BodyText"],
        fontName=font_choice.regular,
        fontSize=9,
        leading=12,
    )
    header_cell_style = ParagraphStyle(
        "MarkToWinTableHeader",
        parent=cell_style,
        fontName=font_choice.bold,
        textColor=colors.HexColor("#0F172A"),
    )

    ensure_parent(output)
    document = SimpleDocTemplate(
        str(output),
        pagesize=A4,
        rightMargin=42,
        leftMargin=42,
        topMargin=48,
        bottomMargin=48,
    )

    story = []
    for block in blocks:
        if block.kind == "heading":
            style = styles["Heading1"] if block.level <= 1 else styles["Heading2"]
            story.extend([Paragraph(paragraph_text(block.text), style), Spacer(1, 6)])
        elif block.kind == "quote":
            story.extend([Paragraph(paragraph_text(block.text), quote_style), Spacer(1, 8)])
        elif block.kind in {"paragraph", "bullet", "number"}:
            prefix = "• " if block.kind == "bullet" else ""
            story.extend([Paragraph(prefix + paragraph_text(block.text), styles["BodyText"]), Spacer(1, 6)])
        elif block.kind == "code":
            story.extend([Preformatted(block.text, code_style), Spacer(1, 8)])
        elif block.kind == "table" and block.rows:
            rows = normalize_table_rows(block.rows)
            table_data = [
                [
                    Paragraph(paragraph_text(cell), header_cell_style if row_index == 0 else cell_style)
                    for cell in row
                ]
                for row_index, row in enumerate(rows)
            ]
            table = Table(table_data, colWidths=table_column_widths(rows, document.width), repeatRows=1)
            table.setStyle(
                TableStyle(
                    [
                        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#E8F1F8")),
                        ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#AAB7C4")),
                        ("FONTNAME", (0, 0), (-1, -1), font_choice.regular),
                        ("FONTNAME", (0, 0), (-1, 0), font_choice.bold),
                        ("LEFTPADDING", (0, 0), (-1, -1), 7),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 7),
                        ("TOPPADDING", (0, 0), (-1, -1), 6),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
                        ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ]
                )
            )
            story.extend([table, Spacer(1, 10)])

    document.build(story)
    return font_choice


def preferred_pdf_font_name(font_path: str | None = None) -> str:
    explicit_name = os.environ.get("MARK_TO_WIN_PDF_CJK_FONT_NAME")
    if explicit_name:
        return explicit_name
    if font_path:
        stem = Path(font_path).stem.lower()
        if stem in {"msyh", "msyhbd", "msyhl"}:
            return "Microsoft YaHei"
        if stem.startswith("noto"):
            return "Noto Sans CJK SC"
        if stem.startswith("simhei"):
            return "SimHei"
        if stem.startswith("simsun"):
            return "SimSun"
        if stem.startswith("deng"):
            return "DengXian"
    if os.name == "nt":
        return "Microsoft YaHei"
    if sys.platform == "darwin":
        return "PingFang SC"
    return "Noto Sans CJK SC"


def select_pdf_engine(requested_engine: str | None = None) -> str | None:
    if requested_engine:
        return requested_engine if shutil.which(requested_engine) else None
    for engine in ("xelatex", "lualatex", "typst", "tectonic", "pdflatex", "wkhtmltopdf", "weasyprint", "prince", "pagedjs-cli"):
        if shutil.which(engine):
            return engine
    return None


def convert_with_pandoc(
    input_path: Path,
    target: str,
    output: Path,
    pdf_engine: str | None = None,
    pdf_font_path: str | None = None,
    pdf_cjk_font_name: str | None = None,
) -> tuple[bool, str]:
    pandoc = shutil.which("pandoc")
    if not pandoc:
        return False, "pandoc is not installed or not on PATH"

    cmd = [
        pandoc,
        str(input_path),
        "--from",
        PANDOC_MARKDOWN_FORMAT,
        "--standalone",
        "--output",
        str(output),
    ]

    if target == "docx":
        pass
    elif target == "pdf":
        engine = select_pdf_engine(pdf_engine)
        if not engine:
            return False, "pandoc PDF output needs a PDF engine on PATH, such as xelatex, lualatex, typst, tectonic, pdflatex, wkhtmltopdf, weasyprint, prince, or pagedjs-cli"
        cmd.extend(["--pdf-engine", engine])
        if engine in {"xelatex", "lualatex"}:
            cjk_font = pdf_cjk_font_name or preferred_pdf_font_name(pdf_font_path)
            cmd.extend(["-V", f"mainfont={cjk_font}", "-V", f"CJKmainfont={cjk_font}"])
    else:
        return False, f"pandoc engine does not support target: {target}"

    ensure_parent(output)
    completed = subprocess.run(cmd, text=True, capture_output=True)
    if completed.returncode == 0:
        return True, "pandoc"

    message = completed.stderr.strip() or completed.stdout.strip() or "pandoc failed"
    return False, message


def find_chromium_browser(requested_browser: str | None = None) -> str | None:
    explicit = requested_browser or os.environ.get("MARK_TO_WIN_BROWSER")
    if explicit:
        explicit_path = Path(explicit).expanduser()
        if explicit_path.exists():
            return str(explicit_path)
        found = shutil.which(explicit)
        if found:
            return found

    for command in ("chrome", "chrome.exe", "chromium", "chromium.exe", "msedge", "msedge.exe"):
        found = shutil.which(command)
        if found:
            return found

    candidates = [
        r"C:\Program Files\Google\Chrome\Application\chrome.exe",
        r"C:\Program Files (x86)\Google\Chrome\Application\chrome.exe",
        r"C:\Program Files\Microsoft\Edge\Application\msedge.exe",
        r"C:\Program Files (x86)\Microsoft\Edge\Application\msedge.exe",
        "/Applications/Google Chrome.app/Contents/MacOS/Google Chrome",
        "/Applications/Microsoft Edge.app/Contents/MacOS/Microsoft Edge",
        "/usr/bin/google-chrome",
        "/usr/bin/chromium",
        "/usr/bin/chromium-browser",
        "/snap/bin/chromium",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return candidate
    return None


def inject_markdown_reader_style(html: str) -> str:
    additions = (
        '<meta name="viewport" content="width=device-width, initial-scale=1">\n'
        f"<style>\n{MARKDOWN_READER_CSS}\n</style>\n"
    )
    html = re.sub(
        r'<pre([^>]*)class="([^"]*\bmermaid\b[^"]*)"([^>]*)>',
        r'<pre\1class="\2 mermaid-unrendered"\3>',
        html,
    )
    if "</head>" in html:
        return html.replace("</head>", additions + "</head>", 1)
    return "<!doctype html><html><head>" + additions + "</head><body>" + html + "</body></html>"


def render_pandoc_html(input_path: Path, html_output: Path) -> tuple[bool, str]:
    pandoc = shutil.which("pandoc")
    if not pandoc:
        return False, "pandoc is not installed or not on PATH"

    cmd = [
        pandoc,
        str(input_path),
        "--from",
        PANDOC_MARKDOWN_FORMAT,
        "--to",
        "html5",
        "--standalone",
        "--mathml",
        "--highlight-style",
        "tango",
        "--output",
        str(html_output),
    ]
    ensure_parent(html_output)
    completed = subprocess.run(cmd, text=True, capture_output=True)
    if completed.returncode != 0:
        message = completed.stderr.strip() or completed.stdout.strip() or "pandoc html conversion failed"
        return False, message

    html = html_output.read_text(encoding="utf-8")
    html_output.write_text(inject_markdown_reader_style(html), encoding="utf-8")
    return True, "pandoc html"


def print_html_to_pdf(html_path: Path, output: Path, browser_path: str | None = None) -> tuple[bool, str]:
    browser = find_chromium_browser(browser_path)
    if not browser:
        return False, "Chrome/Chromium/Edge was not found"

    ensure_parent(output)
    last_message = "browser PDF printing failed"
    for headless_flag in ("--headless=new", "--headless"):
        cmd = [
            browser,
            headless_flag,
            "--disable-gpu",
            "--no-sandbox",
            "--allow-file-access-from-files",
            "--run-all-compositor-stages-before-draw",
            "--virtual-time-budget=3000",
            f"--print-to-pdf={output}",
            "--print-to-pdf-no-header",
            html_path.resolve().as_uri(),
        ]
        completed = subprocess.run(cmd, capture_output=True, timeout=120)
        if completed.returncode == 0 and output.exists() and output.stat().st_size > 0:
            return True, f"chrome html ({browser})"
        stderr = completed.stderr.decode('utf-8', errors='replace').strip() if completed.stderr else ''
        stdout = completed.stdout.decode('utf-8', errors='replace').strip() if completed.stdout else ''
        last_message = stderr or stdout or last_message
    message = last_message
    return False, message


def convert_pdf_with_html(input_path: Path, output: Path, browser_path: str | None = None, keep_html: bool = False) -> tuple[bool, str]:
    if keep_html:
        html_path = output.with_suffix(".html")
        ok, message = render_pandoc_html(input_path, html_path)
        if not ok:
            return False, message
        ok, message = print_html_to_pdf(html_path, output, browser_path)
        return ok, message

    with tempfile.TemporaryDirectory(prefix="mark-to-win-") as tmpdir:
        html_path = Path(tmpdir) / f"{input_path.stem}.html"
        ok, message = render_pandoc_html(input_path, html_path)
        if not ok:
            return False, message
        ok, message = print_html_to_pdf(html_path, output, browser_path)
        return ok, message


def output_path_for(input_path: Path, target: str, output: str | None, output_dir: str | None) -> Path:
    if output:
        return Path(output)
    directory = Path(output_dir) if output_dir else input_path.parent
    return directory / f"{input_path.stem}.{target}"


def main() -> None:
    parser = argparse.ArgumentParser(description="Convert Markdown into Word, PDF, or Excel.")
    parser.add_argument("input", help="Markdown file to convert.")
    parser.add_argument("--to", required=True, help="Target format: docx, pdf, xlsx, or comma-separated list.")
    parser.add_argument("--output", help="Output file path. Use only with a single target.")
    parser.add_argument("--output-dir", help="Directory for generated files.")
    parser.add_argument("--engine", choices=("auto", "native", "pandoc", "html"), default="auto", help="Conversion engine for docx/pdf. Default: auto.")
    parser.add_argument("--pdf-font", help="Optional path to a TrueType/OpenType font for PDF output.")
    parser.add_argument("--pdf-bold-font", help="Optional path to a bold TrueType/OpenType font for PDF headings and table headers.")
    parser.add_argument("--pdf-engine", help="Pandoc PDF engine, such as xelatex or lualatex.")
    parser.add_argument("--pdf-cjk-font-name", help="Font family name passed to Pandoc for CJK PDF output.")
    parser.add_argument("--browser", help="Chrome/Chromium/Edge executable path for the HTML PDF engine.")
    parser.add_argument("--keep-html", action="store_true", help="When using the HTML PDF engine, keep the intermediate styled HTML next to the PDF.")
    parser.add_argument("--show-features", action="store_true", help="Print detected extended Markdown features.")
    args = parser.parse_args()

    input_path = Path(args.input)
    markdown = input_path.read_text(encoding="utf-8")
    features = detect_markdown_features(markdown)
    if args.show_features:
        if features:
            print("Detected extended Markdown features: " + ", ".join(feature.label for feature in features))
        else:
            print("Detected extended Markdown features: none")
    blocks = parse_markdown(markdown)
    targets = [target.strip().lower().lstrip(".") for target in args.to.split(",") if target.strip()]

    if args.output and len(targets) != 1:
        raise SystemExit("--output can only be used with a single target. Use --output-dir for multiple targets.")

    for target in targets:
        output = output_path_for(input_path, target, args.output, args.output_dir)
        if target == "docx":
            if args.engine == "html":
                raise SystemExit("--engine html only supports PDF output. Use --engine pandoc or --engine auto for DOCX.")
            if args.engine in {"auto", "pandoc"}:
                ok, message = convert_with_pandoc(
                    input_path,
                    target,
                    output,
                    pdf_font_path=args.pdf_font,
                    pdf_cjk_font_name=args.pdf_cjk_font_name,
                )
                if ok:
                    print(f"Wrote {output} (engine: pandoc)")
                    continue
                if args.engine == "pandoc":
                    raise SystemExit(f"Pandoc conversion failed: {message}")
                if features:
                    print(f"Warning: Pandoc unavailable for docx ({message}); using native fallback.", file=sys.stderr, flush=True)
                    warn_about_native_limitations(features, target)
            elif features:
                warn_about_native_limitations(features, target)
            convert_docx(blocks, output)
            print(f"Wrote {output} (engine: native)")
        elif target == "xlsx":
            table_count = convert_xlsx(blocks, output)
            print(f"Wrote {output} ({table_count} table(s))")
        elif target == "pdf":
            if args.engine in {"auto", "html"}:
                ok, message = convert_pdf_with_html(input_path, output, args.browser, args.keep_html)
                if ok:
                    print(f"Wrote {output} (engine: html; renderer: {message})")
                    continue
                if args.engine == "html":
                    raise SystemExit(f"HTML PDF conversion failed: {message}")
                print(f"Warning: HTML PDF engine unavailable ({message}); trying Pandoc PDF engine.", file=sys.stderr, flush=True)

            if args.engine in {"auto", "pandoc"}:
                ok, message = convert_with_pandoc(
                    input_path,
                    target,
                    output,
                    pdf_engine=args.pdf_engine,
                    pdf_font_path=args.pdf_font,
                    pdf_cjk_font_name=args.pdf_cjk_font_name,
                )
                if ok:
                    print(f"Wrote {output} (engine: pandoc)")
                    continue
                if args.engine == "pandoc":
                    raise SystemExit(f"Pandoc conversion failed: {message}")
                if features:
                    print(f"Warning: Pandoc unavailable for pdf ({message}); using native fallback.", file=sys.stderr, flush=True)
                    warn_about_native_limitations(features, target)
            elif features:
                warn_about_native_limitations(features, target)
            font_choice = convert_pdf(blocks, output, args.pdf_font, args.pdf_bold_font)
            print(f"Wrote {output} (engine: native; PDF font: {font_choice.source})")
        else:
            raise SystemExit(f"Unsupported target format: {target}")


if __name__ == "__main__":
    main()
